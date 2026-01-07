#!/usr/bin/env python
"""
Markdown to Word Converter
Converts .md (Markdown) files to .docx (Word) format with proper formatting

Formatting:
- Headers: Arial font (size varies by level)
- Body text: Arial 10pt
- Tables: Borders with 0.2cm left/right padding
"""

import sys
import re
import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# Header/Footer settings
HEADER_TITLE = "OCMS Technical Specification"
HEADER_COLOR = RGBColor(0x80, 0x80, 0x80)  # Gray color for header/footer text

# Heading color (dark blue)
HEADING_COLOR = RGBColor(0x1F, 0x4E, 0x79)


# Font settings
FONT_NAME = 'Arial'
BODY_FONT_SIZE = Pt(10)
CODE_FONT_NAME = 'Consolas'
CODE_FONT_SIZE = Pt(9)

# Header font sizes
HEADER_SIZES = {
    1: Pt(16),  # H1
    2: Pt(14),  # H2
    3: Pt(12),  # H3
    4: Pt(11),  # H4
    5: Pt(10),  # H5
    6: Pt(10),  # H6
}

# Table cell padding (0.2 cm = approximately 0.08 inches)
TABLE_CELL_PADDING = Cm(0.2)


def extract_document_title(content):
    """Extract the document title from the first H1 heading."""
    lines = content.split('\n')
    for line in lines:
        if line.startswith('# '):
            return line[2:].strip()
    return "Technical Document"


def extract_prepared_by(content):
    """Extract the company name from 'Prepared by' section."""
    lines = content.split('\n')
    found_prepared_by = False

    for i, line in enumerate(lines):
        # Look for "**Prepared by**" or "Prepared by"
        if 'prepared by' in line.lower():
            found_prepared_by = True
            continue

        # After finding "Prepared by", get the next non-empty line as company name
        if found_prepared_by and line.strip():
            company = line.strip()
            # Remove any markdown formatting
            company = company.replace('**', '').replace('*', '')
            # Stop at horizontal rule or next section
            if company.startswith('---') or company.startswith('#'):
                break
            return company

    return None


def add_header_footer(doc, doc_title):
    """Add header and footer to the document."""
    # Access the default section
    section = doc.sections[0]

    # === HEADER ===
    header = section.header
    header.is_linked_to_previous = False

    # Clear existing header content
    for paragraph in header.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)

    # Add "OCMS Technical Specification" (bold, centered, gray)
    header_para1 = header.add_paragraph()
    header_para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = header_para1.add_run(HEADER_TITLE)
    run1.font.name = FONT_NAME
    run1.font.size = Pt(12)
    run1.font.bold = True
    run1.font.color.rgb = HEADER_COLOR
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    header_para1.paragraph_format.space_after = Pt(0)

    # Add document title (centered, gray)
    header_para2 = header.add_paragraph()
    header_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = header_para2.add_run(doc_title)
    run2.font.name = FONT_NAME
    run2.font.size = Pt(11)
    run2.font.bold = False
    run2.font.color.rgb = HEADER_COLOR
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    header_para2.paragraph_format.space_after = Pt(6)

    # Add horizontal line below header
    header_para3 = header.add_paragraph()
    header_para3.paragraph_format.space_before = Pt(0)
    header_para3.paragraph_format.space_after = Pt(0)

    # Create bottom border for the paragraph (horizontal line)
    pPr = header_para3._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')  # Line thickness
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '808080')  # Gray color
    pBdr.append(bottom)
    pPr.append(pBdr)

    # === FOOTER ===
    footer = section.footer
    footer.is_linked_to_previous = False

    # Clear existing footer content
    for paragraph in footer.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)

    # Add page number (centered)
    footer_para = footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add PAGE field
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

    run.font.name = FONT_NAME
    run.font.size = Pt(10)


def set_section_vertical_alignment(section, alignment='top'):
    """Set vertical alignment for a section.

    Args:
        section: The document section
        alignment: 'top', 'center', 'bottom', or 'both' (justified)
    """
    sectPr = section._sectPr
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), alignment)
    sectPr.append(vAlign)


def add_section_title_page(doc, title_text, is_first_page=False, prepared_by=None):
    """Add a section title as a centered cover page.

    Creates a new page with the title vertically centered,
    followed by a page break so content starts on next page.

    Args:
        doc: The document
        title_text: The title text to display
        is_first_page: If True, this is the document cover (no section break before)
        prepared_by: If provided, adds "Prepared by" and company name below title
    """
    if not is_first_page:
        # Add a section break (next page) to start fresh
        doc.add_section(WD_SECTION_START.NEW_PAGE)
        new_section = doc.sections[-1]
        # Link header/footer to previous section
        new_section.header.is_linked_to_previous = True
        new_section.footer.is_linked_to_previous = True
    else:
        # First page uses the existing first section (header/footer already set)
        new_section = doc.sections[0]

    # Set vertical alignment to center for this section
    set_section_vertical_alignment(new_section, 'center')

    # Add the section title heading
    p = doc.add_heading(title_text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        set_run_font(run, FONT_NAME, HEADER_SIZES[1], bold=True)
        run.font.color.rgb = HEADING_COLOR

    # Add "Prepared by" section for document cover
    if prepared_by and is_first_page:
        # Add some spacing
        doc.add_paragraph()
        doc.add_paragraph()

        # Add "Prepared by" label
        p_label = doc.add_paragraph()
        p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_label = p_label.add_run("Prepared by")
        set_run_font(run_label, FONT_NAME, BODY_FONT_SIZE, bold=True)

        # Add company name
        p_company = doc.add_paragraph()
        p_company.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_company = p_company.add_run(prepared_by)
        set_run_font(run_company, FONT_NAME, BODY_FONT_SIZE)

    # Add another section break for content to start on new page
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    content_section = doc.sections[-1]

    # Link header/footer and set vertical alignment to top for content
    content_section.header.is_linked_to_previous = True
    content_section.footer.is_linked_to_previous = True
    set_section_vertical_alignment(content_section, 'top')


def set_cell_padding(cell, left=None, right=None, top=None, bottom=None):
    """Set padding for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    if left is not None:
        left_elem = OxmlElement('w:left')
        left_elem.set(qn('w:w'), str(int(left.twips)))
        left_elem.set(qn('w:type'), 'dxa')
        tcMar.append(left_elem)

    if right is not None:
        right_elem = OxmlElement('w:right')
        right_elem.set(qn('w:w'), str(int(right.twips)))
        right_elem.set(qn('w:type'), 'dxa')
        tcMar.append(right_elem)

    if top is not None:
        top_elem = OxmlElement('w:top')
        top_elem.set(qn('w:w'), str(int(top.twips)))
        top_elem.set(qn('w:type'), 'dxa')
        tcMar.append(top_elem)

    if bottom is not None:
        bottom_elem = OxmlElement('w:bottom')
        bottom_elem.set(qn('w:w'), str(int(bottom.twips)))
        bottom_elem.set(qn('w:type'), 'dxa')
        tcMar.append(bottom_elem)

    tcPr.append(tcMar)


def set_table_border(table):
    """Set borders for all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)

    tblPr.append(tblBorders)


def shade_cells(row, color='D9E2F3'):
    """Shade cells in a row with specified color."""
    for cell in row.cells:
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, font_name=FONT_NAME, font_size=BODY_FONT_SIZE, bold=False, italic=False):
    """Set font properties for a run."""
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    run.italic = italic
    # Set font for East Asian and complex scripts as well
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def format_json_string(json_str):
    """Try to parse and pretty-print JSON string."""
    try:
        # Try to parse as JSON
        parsed = json.loads(json_str)
        # Pretty print with 2-space indentation
        return json.dumps(parsed, indent=2, ensure_ascii=False)
    except (json.JSONDecodeError, ValueError):
        # Not valid JSON, try manual formatting for JSON-like strings
        # This handles cases with placeholders like [...] or <...>
        result = json_str
        # Add newline after { and [
        result = re.sub(r'(\{)\s*', '{\n  ', result)
        result = re.sub(r'(\[)\s*(?!")', '[\n    ', result)
        # Add newline before } and ]
        result = re.sub(r'\s*(\})', '\n}', result)
        result = re.sub(r'\s*(\])(?!,|\s*\})', '\n  ]', result)
        # Add newline after commas (for top-level properties)
        result = re.sub(r'",\s*"', '",\n  "', result)
        result = re.sub(r'(\d),\s*"', r'\1,\n  "', result)
        result = re.sub(r'(\]),\s*"', r'],\n  "', result)
        result = re.sub(r'(\}),\s*"', r'},\n  "', result)
        return result


def is_json_like(text):
    """Check if text looks like JSON (starts with { or [)."""
    stripped = text.strip()
    return (stripped.startswith('{') and stripped.endswith('}')) or \
           (stripped.startswith('[') and stripped.endswith(']'))


def add_formatted_text(paragraph, text, font_name=FONT_NAME, font_size=BODY_FONT_SIZE):
    """Add text with inline formatting (bold, italic, code)."""
    # Replace <br> and <br/> tags with newline
    text = re.sub(r'<br\s*/?>', '\n', text)

    # Pattern to match **bold**, *italic*, `code`
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, font_name, font_size, bold=True)
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, font_name, font_size, italic=True)
        elif part.startswith('`') and part.endswith('`'):
            code_content = part[1:-1]
            # Check if it's JSON-like content
            if is_json_like(code_content):
                formatted_json = format_json_string(code_content)
                run = paragraph.add_run(formatted_json)
            else:
                run = paragraph.add_run(code_content)
            # Use Arial 10pt for all content
            set_run_font(run, FONT_NAME, BODY_FONT_SIZE)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, font_name, font_size)


def parse_markdown_table(lines, start_index):
    """Parse a markdown table and return table data and end index."""
    table_data = []
    i = start_index

    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith('|'):
            break

        # Skip separator line (|---|---|)
        if re.match(r'^\|[\s\-:|]+\|$', line):
            i += 1
            continue

        # Parse cells
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells:
            table_data.append(cells)
        i += 1

    return table_data, i


def setup_document_styles(doc):
    """Setup document styles with Arial font."""
    # Set default Normal style
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = BODY_FONT_SIZE
    style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    # Set heading styles
    for level in range(1, 7):
        style_name = f'Heading {level}'
        if style_name in doc.styles:
            heading_style = doc.styles[style_name]
            heading_style.font.name = FONT_NAME
            heading_style.font.size = HEADER_SIZES.get(level, Pt(10))
            heading_style.font.bold = True
            heading_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    # Set list styles
    for style_name in ['List Bullet', 'List Bullet 2', 'List Number']:
        if style_name in doc.styles:
            list_style = doc.styles[style_name]
            list_style.font.name = FONT_NAME
            list_style.font.size = BODY_FONT_SIZE


def convert_markdown_to_word(input_path: str, output_path: str = None) -> str:
    """Convert a Markdown file to Word document format."""
    input_path = Path(input_path)

    if not input_path.exists():
        raise FileNotFoundError(f'File not found: {input_path}')

    if input_path.suffix.lower() != '.md':
        raise ValueError(f'File must be a .md file, got: {input_path.suffix}')

    if output_path is None:
        output_path = input_path.with_suffix('.docx')
    else:
        output_path = Path(output_path)

    # Read markdown content
    with open(input_path, 'r', encoding='utf-8') as md_file:
        content = md_file.read()

    # Create Word document
    doc = Document()

    # Setup document styles
    setup_document_styles(doc)

    # Extract document title and add header/footer
    doc_title = extract_document_title(content)
    add_header_footer(doc, doc_title)

    # Extract "Prepared by" company name for cover page
    prepared_by_company = extract_prepared_by(content)

    # Process markdown content
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_content = []
    first_h1_seen = False  # Track if first H1 has been processed
    skip_prepared_by = False  # Flag to skip "Prepared by" section after cover

    while i < len(lines):
        line = lines[i]

        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_block_content)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                # Use Arial 10pt for code blocks
                set_run_font(run, FONT_NAME, BODY_FONT_SIZE)
                p.paragraph_format.left_indent = Inches(0.5)
                code_block_content = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue

        # Skip HTML comments
        if line.strip().startswith('<!--'):
            while i < len(lines) and '-->' not in lines[i]:
                i += 1
            i += 1
            continue

        # Skip "Prepared by" section after cover page (it's already on cover)
        if skip_prepared_by:
            stripped = line.strip().lower()
            # Skip until we hit horizontal rule or next section
            if stripped.startswith('---') or stripped.startswith('#'):
                skip_prepared_by = False
                # Don't skip the current line if it's a heading or rule
                if not stripped.startswith('#'):
                    i += 1
                    continue
            else:
                i += 1
                continue

        # Handle headers
        if line.startswith('# '):
            heading_text = line[2:].strip()

            if first_h1_seen:
                # Section titles (not doc title) get their own centered cover page
                add_section_title_page(doc, heading_text)
            else:
                # First H1 is document title - also gets centered cover page
                add_section_title_page(doc, heading_text, is_first_page=True, prepared_by=prepared_by_company)
                # Skip the "Prepared by" section in the content since it's now on cover
                if prepared_by_company:
                    skip_prepared_by = True

            first_h1_seen = True
            i += 1
            continue
        elif line.startswith('## '):
            heading_text = line[3:].strip()
            # Add page break before Table of Contents (Version History already on new page after cover)
            if heading_text.lower() == 'table of contents':
                doc.add_page_break()
            p = doc.add_heading(heading_text, level=2)
            for run in p.runs:
                set_run_font(run, FONT_NAME, HEADER_SIZES[2], bold=True)
                run.font.color.rgb = HEADING_COLOR
            i += 1
            continue
        elif line.startswith('### '):
            p = doc.add_heading(line[4:].strip(), level=3)
            for run in p.runs:
                set_run_font(run, FONT_NAME, HEADER_SIZES[3], bold=True)
                run.font.color.rgb = HEADING_COLOR
            i += 1
            continue
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:].strip(), level=4)
            for run in p.runs:
                set_run_font(run, FONT_NAME, HEADER_SIZES[4], bold=True)
                run.font.color.rgb = HEADING_COLOR
            i += 1
            continue
        elif line.startswith('##### '):
            p = doc.add_heading(line[6:].strip(), level=5)
            for run in p.runs:
                set_run_font(run, FONT_NAME, HEADER_SIZES[5], bold=True)
                run.font.color.rgb = HEADING_COLOR
            i += 1
            continue
        elif line.startswith('###### '):
            p = doc.add_heading(line[7:].strip(), level=6)
            for run in p.runs:
                set_run_font(run, FONT_NAME, HEADER_SIZES[6], bold=True)
                run.font.color.rgb = HEADING_COLOR
            i += 1
            continue

        # Handle horizontal rule
        if line.strip() in ['---', '***', '___']:
            # Add a thin horizontal line
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            i += 1
            continue

        # Handle tables
        if line.strip().startswith('|'):
            table_data, end_index = parse_markdown_table(lines, i)
            if table_data:
                # Create table
                num_cols = len(table_data[0])
                table = doc.add_table(rows=len(table_data), cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
                set_table_border(table)

                for row_idx, row_data in enumerate(table_data):
                    row = table.rows[row_idx]
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < len(row.cells):
                            cell = row.cells[col_idx]

                            # Set cell padding (0.2 cm left and right)
                            set_cell_padding(cell, left=TABLE_CELL_PADDING, right=TABLE_CELL_PADDING)

                            # Clear existing content and add formatted text
                            cell.text = ''
                            paragraph = cell.paragraphs[0]

                            # Use add_formatted_text for proper inline formatting (bold, italic, code/JSON)
                            if row_idx == 0:
                                # Header row - bold text
                                run = paragraph.add_run(cell_text)
                                set_run_font(run, FONT_NAME, BODY_FONT_SIZE, bold=True)
                            else:
                                # Data rows - use formatted text to handle JSON, code, etc.
                                add_formatted_text(paragraph, cell_text, FONT_NAME, BODY_FONT_SIZE)

                    # Shade header row
                    if row_idx == 0:
                        shade_cells(row, 'D9E2F3')

                doc.add_paragraph()  # Add space after table
            i = end_index
            continue

        # Handle bullet lists
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text, FONT_NAME, BODY_FONT_SIZE)
            i += 1
            continue

        # Handle numbered lists (use manual numbering to avoid Word auto-continue)
        numbered_match = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
        if numbered_match:
            number = numbered_match.group(1)
            text = numbered_match.group(2)
            p = doc.add_paragraph()
            # Add number manually
            run_num = p.add_run(f"{number}. ")
            set_run_font(run_num, FONT_NAME, BODY_FONT_SIZE)
            # Add the rest of the text
            add_formatted_text(p, text, FONT_NAME, BODY_FONT_SIZE)
            # Add left indent for numbered list appearance
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            i += 1
            continue

        # Handle indented bullet lists (sub-items)
        if line.startswith('  - ') or line.startswith('  * '):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet 2')
            add_formatted_text(p, text, FONT_NAME, BODY_FONT_SIZE)
            i += 1
            continue

        # Handle images
        img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line.strip())
        if img_match:
            alt_text = img_match.group(1)
            img_path = img_match.group(2)

            # Try to add image if it exists
            full_img_path = input_path.parent / img_path
            if full_img_path.exists():
                try:
                    doc.add_picture(str(full_img_path), width=Inches(6))
                except Exception as e:
                    p = doc.add_paragraph()
                    run = p.add_run(f'[Image: {alt_text}] ({img_path})')
                    set_run_font(run, FONT_NAME, BODY_FONT_SIZE, italic=True)
            else:
                p = doc.add_paragraph()
                run = p.add_run(f'[Image: {alt_text}] ({img_path})')
                set_run_font(run, FONT_NAME, BODY_FONT_SIZE, italic=True)
            i += 1
            continue

        # Handle NOTE: lines
        if line.strip().startswith('NOTE:'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip())
            set_run_font(run, FONT_NAME, BODY_FONT_SIZE, italic=True)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)
            i += 1
            continue

        # Handle bold lines (like **Note:**)
        if line.strip().startswith('**') and '**' in line.strip()[2:]:
            p = doc.add_paragraph()
            add_formatted_text(p, line.strip(), FONT_NAME, BODY_FONT_SIZE)
            i += 1
            continue

        # Handle regular paragraphs
        stripped = line.strip()
        if stripped:
            p = doc.add_paragraph()
            add_formatted_text(p, stripped, FONT_NAME, BODY_FONT_SIZE)

        i += 1

    # Save document
    doc.save(str(output_path))
    return str(output_path)


def convert_directory(input_dir: str, output_dir: str = None) -> list:
    """Convert all Markdown files in a directory to Word documents."""
    input_dir = Path(input_dir)
    output_dir = Path(output_dir) if output_dir else input_dir

    if not input_dir.is_dir():
        raise NotADirectoryError(f'Not a directory: {input_dir}')

    output_dir.mkdir(parents=True, exist_ok=True)
    converted_files = []
    md_files = list(input_dir.glob('*.md'))

    if not md_files:
        print(f'No .md files found in {input_dir}')
        return converted_files

    for md_file in md_files:
        output_path = output_dir / md_file.with_suffix('.docx').name
        try:
            result = convert_markdown_to_word(str(md_file), str(output_path))
            converted_files.append(result)
            print(f'Converted: {md_file.name} -> {output_path.name}')
        except Exception as e:
            print(f'Error converting {md_file.name}: {e}')

    return converted_files


def main():
    if len(sys.argv) < 2:
        print('Markdown to Word Converter')
        print('-' * 40)
        print('Usage:')
        print('  Single file:  python md_to_word.py <input.md> [output.docx]')
        print('  Directory:    python md_to_word.py <input_dir> [output_dir] --dir')
        print()
        print('Examples:')
        print('  python md_to_word.py document.md')
        print('  python md_to_word.py document.md output.docx')
        print('  python md_to_word.py ./docs --dir')
        print('  python md_to_word.py ./docs ./word_output --dir')
        print()
        print('Formatting:')
        print(f'  Font: {FONT_NAME}')
        print(f'  Body text: {BODY_FONT_SIZE.pt}pt')
        print(f'  Headers: {HEADER_SIZES[1].pt}pt (H1) to {HEADER_SIZES[6].pt}pt (H6)')
        print(f'  Table padding: {TABLE_CELL_PADDING.cm}cm left/right')
        print()
        print('Requirements:')
        print('  pip install python-docx')
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 and sys.argv[2] != '--dir' else None
    is_directory = '--dir' in sys.argv

    try:
        if is_directory:
            results = convert_directory(input_path, output_path)
            print(f'\nConverted {len(results)} file(s)')
        else:
            result = convert_markdown_to_word(input_path, output_path)
            print(f'Successfully converted to: {result}')
    except Exception as e:
        print(f'Error: {e}')
        sys.exit(1)


if __name__ == '__main__':
    main()
